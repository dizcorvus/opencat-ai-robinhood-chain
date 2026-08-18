import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(17)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
    elif level == 2:
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(71, 85, 105)
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(3)
    return h

def add_callout(doc, text, title="💡 CATATAN PENTING", border_color="4F46E5", bg_color="F8FAFC"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r_title = p.add_run(f"{title}\n")
    r_title.font.name = 'Calibri'
    r_title.font.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(
        int(border_color[0:2], 16),
        int(border_color[2:4], 16),
        int(border_color[4:6], 16)
    )
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(9.5)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_code_block(doc, code_text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "0F172A")
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(241, 245, 249)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E293B")
        set_cell_margins(hdr_cells[i], top=80, bottom=80, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx+1].cells
        bg_color = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=100, right=100)
            p = row_cells[c_idx].paragraphs[0]
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(51, 65, 85)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_paragraph_styled(doc, text, bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Calibri'
        r_bold.font.bold = True
        r_bold.font.size = Pt(10)
        r_bold.font.color.rgb = RGBColor(15, 23, 42)
        
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_bullet_styled(doc, title, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    r_title = p.add_run(f"{title}: ")
    r_title.font.name = 'Calibri'
    r_title.font.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(15, 23, 42)
    
    r_text = p.add_run(text)
    r_text.font.name = 'Calibri'
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    return p
