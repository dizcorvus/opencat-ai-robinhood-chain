import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
sys.path.insert(0, SCRIPT_DIR)

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stderr, 'reconfigure'):
    sys.stderr.reconfigure(encoding='utf-8')

from builder_core import (
    add_styled_heading,
    add_callout,
    add_code_block,
    add_styled_table,
    add_paragraph_styled,
    add_bullet_styled
)

ID_MARKDOWN = """# 🐾 Mengenal Opencatz AI: Pasukan AI Agent Otonom Pemburu Alpha di Robinhood Chain

> *"Chill trades, 9 lives, razor-sharp on-chain alpha."* — **Opencatz AI** 🐾⚡
>
> 🌐 **Website Resmi:** [https://opencatz.xyz](https://opencatz.xyz)  
> 📖 **Dokumentasi:** [https://opencatz.xyz/docs](https://opencatz.xyz/docs)  
> 💻 **Web Terminal Emulator:** [https://opencatz.xyz/terminal](https://opencatz.xyz/terminal)  
> 🔗 **GitHub Repository:** [https://github.com/dizcorvus/opencatz-ai-robinhood-chain](https://github.com/dizcorvus/opencatz-ai-robinhood-chain)  
>
> ⚠️ **Disclaimer Awal (NFA & DYOR):** Seluruh konten dalam artikel ini ditujukan untuk edukasi, eksplorasi teknologi Web3, dan riset kuantitatif. Trading aset kripto on-chain, meme coin, liquidity pool, dan NFT memiliki risiko volatilitas tinggi. Selalu lakukan riset mandiri (*Do Your Own Research*), gunakan manajemen risiko yang terukur, dan jangan pernah menggunakan dana darurat. *Not Financial Advice.*

---

## 1. Masalah Nyata On-Chain: Kenapa Manual Screening Bikin Capek?

Kalau kamu pernah aktif trading di jaringan Layer-2 baru, kamu pasti paham situasinya: ratusan token baru lahir setiap jam, linimasa X (Twitter) langsung dibanjiri narasi koin baru, dan chart harga bergerak sangat liar.

Peluang profit di ekosistem baru memang sangat besar, namun realita trading on-chain memiliki tantangan berat:
* **Banjir Koin Scam & Rugpull:** Sekitar 90% token baru yang baru dideploy biasanya berakhir jadi honeypot, likuiditas ditarik dev (*rugpull*), atau langsung dihajar sniper bot.
* **Keterbatasan Fisik Manusia:** Pasar crypto bergerak 24/7 tanpa henti. Sering kali momentum rotasi likuiditas besar atau alpha terbaik justru terjadi jam 3 pagi saat kita sedang tidur.
* **Fragmentasi Data & Tab Berantakan:** Mau riset satu token saja harus membuka DexScreener (chart), GoPlus (audit kontrak), GMGN (pelacakan smart money), Twitter/X (sentimen), dan Krystal (pool likuiditas). Sangat melelahkan dan memakan waktu.
* **Trading Emosional & FOMO:** Melihat candle hijau panjang langsung mendorong beli di puncak, lalu panik jual saat terjadi koreksi normal.

Dari masalah nyata inilah **Opencatz AI** diciptakan. Opencatz AI bukan sekadar bot alert webhook biasa, melainkan ekosistem **Multi-Agent Swarm Intelligence** otonom yang memantau Robinhood Chain 24/7, menyaring noise lewat konsensus berlapis, dan mengirimkan sinyal trading siap eksekusi ke Discord, Terminal TUI, dan Telegram.

---

## 2. Jaringan Spesialis: Robinhood Chain (EVM L2 #4663)

Opencatz AI tidak dibuat campur aduk dengan bridge lintas-rantai yang lambat dan rawan eksploit. Sistem ini dirancang native dan fokus 100% pada **Robinhood Chain**.

| Parameter Jaringan | Spesifikasi |
| :--- | :--- |
| **Nama Jaringan** | Robinhood Chain (EVM Layer 2) |
| **Chain ID** | `4663` |
| **Native Asset** | `ETH` |
| **Canonical RPC** | `https://rpc.mainnet.chain.robinhood.com` |
| **Block Explorer** | `https://robinhoodchain.blockscout.com` |
| **Primary DEX Venue** | Uniswap V3 Router (Robinhood Chain L2) |
| **Karakteristik** | Eksekusi sub-detik dengan gas fee sangat murah |

Fokus pada satu rantai tunggal memastikan latensi super rendah, keandalan eksekusi swap, dan bebas dari risiko jembatan bridge.

---

## 3. Bedah Arsitektur Swarm: Cara Kerja Multi-Agent Opencatz

Opencatz AI mendistribusikan beban kerja ke 5 sub-agent screening spesialis, 1 mesin konsensus swarm, dan modul manajemen risiko otomatis:

```
                          USER INTERFACE PLATFORMS
              (Discord Command Center · Terminal TUI · Telegram Bridge)
                                      │
                                      ▼
                   ┌───────────────────────────────────┐
                   │       OPENCATZ CORE HUB           │
                   │   #opencatz-control-room · chat   │
                   │   risk gate · 9-lives risk engine │
                   │   wallet service · trade journal  │
                   └──────────────────┬────────────────┘
                                      │ candidate signals
     ┌────────────────┬───────────────┼───────────────┬────────────────┐
     ▼                ▼               ▼               ▼                ▼
 ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────┐
 │  MEME AGENT  │ │ LP VELOCITY  │ │  NFT AGENT   │ │ ALPHA SCRAPER│ │ ETH WHALES   │
 │meme-robinhood│ │ lp-robinhood │ │    nft       │ │alpha-robinhood│ │ whale-eth    │
 │ GMGN + GoPlus│ │ Krystal Cloud│ │   OpenSea    │ │X API v2 / Web│ │ Hyperliquid  │
 │vol 24h ≥ $25k│ │Fee/TVL ≥ 2%  │ │floor +10%/1h │ │1h rh alpha   │ │perps ≥ $500k │
 │liq ≥ $5k     │ │TVL ≥ $10k    │ │sales ≥ 3/h   │ │sentiment high│ │spot ≥ $50k   │
 └──────┬───────┘ └──────┬───────┘ └──────┬───────┘ └──────┬───────┘ └──────┬───────┘
        └────────────────┴───────────────┼────────────────┴────────────────┘
                                         ▼
                   ┌────────────────────────────────────┐
                   │   SWARM CONSENSUS ENGINE (≥ 80%)   │
                   │  Quant · Catalyst · Security Audit │
                   └────────────────┬───────────────────┘
                                    │ only ≥ 80% confidence
                                    ▼
          MULTI-PLATFORM DISPATCH (Discord · Terminal TUI · Telegram)
                                    │
                                    ▼
                        WALLET TRACKER & POSITION MANAGER
```

### Lima Sub-Agent Screening Spesialis
1. 🌸 **Meme Robinhood Agent (`#call-meme-robinhood`):** Berburu token meme potensial via GMGN OpenAPI & GoPlus Security. Filter bawaan: volume 24 jam minimal $25k, likuiditas minimal $5k, fee pool minimal $250, dan wajib lolos audit keamanan kontrak.
2. 🌊 **LP Velocity Agent (`#call-lp-robinhood`):** Memantau Concentrated Liquidity Uniswap V3 via Krystal Cloud API (`ethereum@4663`) dengan kriteria TVL ≥ $10k, volume 24 jam ≥ $100k, dan rasio Fee/TVL ≥ 2% untuk passive income.
3. 🔮 **NFT Sniper Agent (`#call-nft-robinhood`):** Memantau pasar NFT seperti Catz NFT di OpenSea REST API v2 dengan filter kenaikan floor ≥ +10%/jam, lonjakan volume ≥ 1.5x, dan sales velocity minimal 3 transaksi/jam.
4. ☀️ **Alpha Scraper & Sentimen X (`#call-alpha-robinhood`):** Mengikis narasi koin baru di Robinhood Chain dalam 1 jam terakhir dan mengevaluasi sentimen sosial via official X API v2.
5. 🐋 **ETH Whale Tracker (`#call-whale-eth`):** Melacak pergerakan modal besar institusional di Hyperliquid L1 (posisi perpetual ETH ≥ $500k dan order flow spot ≥ $50k).

### 3-Layer Swarm Consensus Engine (Minimal 80%)
Tidak ada satu agen pun yang dapat memposting sinyal tanpa persetujuan swarm. Setiap kandidat token diuji lewat 3 lapisan filter:
* **Layer Kuantitatif & Likuiditas (35%):** Kedalaman pool, slippage, rasio volume/market cap.
* **Layer Katalis & Sentimen (35%):** Volume velocity, sentimen sosial, akumulasi smart money.
* **Layer Audit Keamanan (30%):** Deteksi honeypot, mintability, blacklist, dan kepemilikan dev (*fail-closed*).

Jika skor gabungan kurang dari **80%**, sinyal otomatis ditolak. Dilengkapi juga dengan **Cross-Agent Conflict Veto**: jika agen whale mendeteksi sinyal SHORT atau dumping pada aset yang sama, sistem langsung membatalkan rekomendasi BUY dari agen lain.

### 9-Lives Risk Engine & Position Manager
Setelah posisi terbuka, Position Manager otomatis mengawal:
* **Circuit Breaker:** Pengaman darurat yang menghentikan bot otomatis saat drawdown pasar harian mencapai batas toleransi.
* **Take Profit (TP) Bertingkat:** Peringatan profit otomatis di target +100% (2x) dan +200% (3x).
* **Stop Loss (SL) Disiplin:** Batas proteksi modal tegas di angka -20%.
* **Dynamic Trailing Stop:** Mengunci profit yang telah diperoleh saat harga terus naik.
* **Peringatan LP Out-of-Range:** Notifikasi instan saat harga keluar dari rentang konsentrasi likuiditas Uniswap V3.

---

## 4. Kustomisasi Penuh & Custom Strategy Compiler

Opencatz AI memberikan kebebasan penuh bagi trader untuk merancang gaya trading masing-masing:
* **Kompilasi Bahasa Manusia (Natural Language Prompt):** Kamu cukup menulis aturan screening menggunakan bahasa sehari-hari (misal: *"Hanya cari token meme yang dipegang minimal 3 smart wallet dan likuiditas di atas $15k"*). Saat bot pertama kali booting, prompt ini otomatis dikompilasi menjadi modul strategi sandboxed `.mjs` yang aman.
* **Preset Strategi Siap Pakai:**
  * *Loosened Default:* Menghasilkan sinyal ~2x lebih aktif untuk trader yang menyukai frekuensi tinggi, dengan tetap menjaga batas kualitas 80%.
  * *Standard:* Filter ketat dan konservatif.
  * *Numeric Editor:* Mengubah parameter angka secara langsung lewat antarmuka onboarding atau Discord chat.
* **Folder Indikator Kustom (`indicators/`):** Tempat menyematkan formula indikator teknikal kustom buatanmu sendiri.

---

## 5. 💡 Tips Hemat: Jalankan 100% Gratis Pakai OpenRouter Free Tier

Kamu bisa menjalankan Opencatz AI secara penuh **tanpa biaya langganan AI sepeser pun ($0/bulan)**:
1. Buat akun di [OpenRouter.ai](https://openrouter.ai) dan dapatkan API key gratis.
2. Saat onboarding, pilih provider `openrouter` dan gunakan model gratis berkualitas tinggi (seperti `meta-llama/llama-3.3-70b-instruct:free` atau `deepseek/deepseek-r1:free`).
3. Opencatz AI dirancang sangat efisien. Seluruh komputasi berat, filtering data, dan audit keamanan dijalankan menggunakan kalkulasi matematis lokal. LLM hanya dipanggil untuk tugas penalaran tingkat tinggi (sentimen tweet dan interaksi chat control room). Biaya operasional bot benar-benar **$0**!

---

## 6. Tiga Mode Eksekusi & Pusat Kendali Multi-Platform

### Mode Eksekusi
1. **`DRY_RUN` (Default):** Simulasi pasar 100% realistis dengan harga live Uniswap V3 tanpa memotong saldo dompet asli. Sangat aman untuk pemula dan pengujian strategi.
2. **`SIGNAL_ONLY`:** Bot bekerja sebagai radar intelijen murni yang membagikan call card lengkap dengan tautan swap untuk eksekusi manual.
3. **`AUTO_EXECUTE`:** Bot mengeksekusi swap otomatis di on-chain via Viem saat sinyal mencapai konsensus ≥ 80% dan lolos seluruh aturan manajemen risiko.

### Antarmuka Multi-Platform
* **🎮 Discord Command Center:** Kategori `🐾 OPENCATZ COMMAND CENTER` dengan 6 channel otomatis dan 22 slash command.
* **💻 Terminal TUI (`opencatz terminal`):** Tampilan konsol 24-bit TrueColor ANSI yang responsif untuk VPS Linux.
* **📱 Telegram Notification Bridge:** Notifikasi push real-time ke smartphone dengan tombol aksi cepat.
* **🌐 Web Dashboard REST API (Port 3000):** Siap dihubungkan ke dashboard web Next.js atau aplikasi mobile kustom.

---

## 7. 🎟️ Live Deployment di PX Identities Discord (Khusus Holders 404 Identities)

Bagi kamu yang ingin langsung menikmati sinyal tanpa perlu setup server atau VPS sendiri: **Opencatz AI akan segera hadir secara live 24/7 di server Discord PX Identities!**

Seluruh holder koleksi **404 Identities (Robinhood Chain)** akan mendapatkan akses eksklusif:
* Akses penuh ke seluruh channel screening sinyal real-time.
* Audit smart contract on-demand langsung di Discord.
* Notifikasi pergerakan whale dan rangkuman alpha harian tanpa biaya server tambahan.

---

## 8. Panduan Instalasi Langkah Demi Langkah

### Prasyarat
* **Node.js** versi 22.12 ke atas ([Unduh Node.js](https://nodejs.org))
* **Git**
* **Discord Bot Token** dari Discord Developer Portal

### Langkah 1: Clone Repositori
```bash
git clone https://github.com/dizcorvus/opencatz-ai-robinhood-chain.git
cd "Opencatz AI (Robinhood Chain)"
```

### Langkah 2: Setup Otomatis 1-Klik
* **Pengguna Windows (PowerShell):**
  ```powershell
  .\\setup.bat
  ```
* **Pengguna Linux / macOS / VPS:**
  ```bash
  bash setup.sh
  ```

### Langkah 3: Konfigurasi Interaktif (`opencatz onboard`)
Jalankan wizard konfigurasi:
```bash
opencatz onboard
```
Wizard akan memandu pengaturan mode eksekusi, token Discord/Telegram, AI provider, serta API key data (GMGN, Krystal Cloud, OpenSea, GoPlus, Uniswap, X API v2) beserta kunci cadangannya (`*_BACKUP_KEYS`).

### Langkah 4: Jalankan Bot
* **Mode Development:** `opencatz run`
* **Mode Terminal TUI:** `opencatz terminal`
* **Mode Produksi 24/7 (PM2 Daemon):** `opencatz deploy`

> ✨ Saat pertama kali bot online di server Discord, Opencatz akan **otomatis membuat kategori `🐾 OPENCATZ COMMAND CENTER`, 6 channel, dan mendaftarkan seluruh 22 slash command**.

### Perintah Penting
* `/analyze [address]`: Audit instan 3 lapis untuk sembarang kontrak token.
* `/wallet balance / setup`: Kelola burner wallet dan cek saldo on-chain.
* `/alert set [token] [target]`: Pasang alarm target harga otomatis.
* `/journal summary / export`: Rekap win-rate dan ekspor riwayat trading ke CSV.
* `opencatz doctor`: Diagnostik menyeluruh koneksi RPC, API key, dan status sub-agent.
* `opencatz update`: Pembaruan otomatis 1-perintah (git pull, build, restart PM2).

---

## 9. 🌟 Open Source & Bocoran Edisi Mendatang

Opencatz AI adalah proyek **100% Open-Source** di bawah lisensi MIT. Kami mengundang seluruh developer, trader on-chain, dan peneliti AI untuk berkolaborasi dan berkontribusi.

🔗 **GitHub Repository:** [https://github.com/dizcorvus/opencatz-ai-robinhood-chain](https://github.com/dizcorvus/opencatz-ai-robinhood-chain)

Sebagai penutup, arsitektur di Robinhood Chain ini adalah pondasi awal. Tim sedang mengembangkan **Opencatz AI Multi-Chain Edition** (Solana, Base, Arbitrum, BSC) serta **Premium Swarm Execution Engine**. Ikuti terus pembaruannya di [opencatz.xyz](https://opencatz.xyz)!
"""

EN_MARKDOWN = """# 🐾 Inside Opencatz AI: The Autonomous Multi-Agent Trading Swarm on Robinhood Chain

> *"Chill trades, 9 lives, razor-sharp on-chain alpha."* — **Opencatz AI** 🐾⚡
>
> 🌐 **Official Website:** [https://opencatz.xyz](https://opencatz.xyz)  
> 📖 **Documentation:** [https://opencatz.xyz/docs](https://opencatz.xyz/docs)  
> 💻 **Web Terminal Emulator:** [https://opencatz.xyz/terminal](https://opencatz.xyz/terminal)  
> 🔗 **GitHub Repository:** [https://github.com/dizcorvus/opencatz-ai-robinhood-chain](https://github.com/dizcorvus/opencatz-ai-robinhood-chain)  
>
> ⚠️ **Upfront Disclaimer (NFA & DYOR):** This article is strictly for educational, research, and technical exploratory purposes. Trading on-chain crypto assets, meme tokens, liquidity pools, and NFTs carries substantial financial risk. Always Do Your Own Research (DYOR) and employ strict capital risk management. *Not Financial Advice.*

---

## 1. The On-Chain Reality: Why Manual Screening Falls Short

Trading on a newly launched Layer-2 network is exhilarating, but anyone active on-chain understands the harsh reality: hundreds of new contracts deploy every hour, social feeds are saturated with promotional noise, and liquidity moves at sub-second velocity.

While the upside in new ecosystems is immense, individual traders face steep obstacles:
* **Overwhelming Scam Influx & Rugpulls:** Approximately 90% of newly deployed tokens end up as honeypots, developer liquidity removals (rugpulls), or illiquid pools dominated by automated sniper bots.
* **Human Physical Fatigue:** Crypto markets operate 24/7. High-conviction alpha and critical liquidity migrations frequently occur at 3:00 AM while you are asleep.
* **Severe Tool Fragmentation:** Vetting a single token requires juggling DexScreener (charts), GoPlus (contract security), GMGN (smart money flows), Twitter/X (sentiment), and Krystal (pool yields). It is exhausting and slow.
* **Emotional Traps & FOMO:** Seeing sudden green candles tempts traders to buy at local tops, only to panic sell during standard market corrections.

**Opencatz AI** was engineered to solve these exact structural problems. Rather than acting as a simple alert webhook, it is a comprehensive **Multi-Agent Swarm Intelligence** that autonomously monitors Robinhood Chain 24/7, filters noise through strict multi-agent consensus, and delivers actionable intelligence to Discord, an interactive terminal console, or Telegram.

---

## 2. Dedicated Single-Chain Focus: Robinhood Chain (EVM L2 #4663)

Opencatz AI avoids fragile cross-chain bridges and complex routing layers by operating natively and exclusively on **Robinhood Chain**.

| Network Parameter | Specification |
| :--- | :--- |
| **Network Name** | Robinhood Chain (EVM Layer 2) |
| **Chain ID** | `4663` |
| **Native Asset** | `ETH` |
| **Canonical RPC** | `https://rpc.mainnet.chain.robinhood.com` |
| **Block Explorer** | `https://robinhoodchain.blockscout.com` |
| **Primary DEX Venue** | Uniswap V3 Router (Robinhood Chain L2) |
| **Execution Profile** | Sub-second transaction finality with negligible gas fees |

Focusing on a dedicated single-chain stack guarantees ultra-low latency, reliable swap execution, and zero bridge-related vulnerability vectors.

---

## 3. System Architecture: How the Swarm Operates

Opencatz AI divides screening and trading responsibilities across 5 specialist screening agents, a central consensus gate, and post-execution risk managers:

```
                          USER INTERFACE PLATFORMS
              (Discord Command Center · Terminal TUI · Telegram Bridge)
                                      │
                                      ▼
                   ┌───────────────────────────────────┐
                   │       OPENCATZ CORE HUB           │
                   │   #opencatz-control-room · chat   │
                   │   risk gate · 9-lives risk engine │
                   │   wallet service · trade journal  │
                   └──────────────────┬────────────────┘
                                      │ candidate signals
     ┌────────────────┬───────────────┼───────────────┬────────────────┐
     ▼                ▼               ▼               ▼                ▼
 ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────┐ ┌──────────────┐
 │  MEME AGENT  │ │ LP VELOCITY  │ │  NFT AGENT   │ │ ALPHA SCRAPER│ │ ETH WHALES   │
 │meme-robinhood│ │ lp-robinhood │ │    nft       │ │alpha-robinhood│ │ whale-eth    │
 │ GMGN + GoPlus│ │ Krystal Cloud│ │   OpenSea    │ │X API v2 / Web│ │ Hyperliquid  │
 │vol 24h ≥ $25k│ │Fee/TVL ≥ 2%  │ │floor +10%/1h │ │1h rh alpha   │ │perps ≥ $500k │
 │liq ≥ $5k     │ │TVL ≥ $10k    │ │sales ≥ 3/h   │ │sentiment high│ │spot ≥ $50k   │
 └──────┬───────┘ └──────┬───────┘ └──────┬───────┘ └──────┬───────┘ └──────┬───────┘
        └────────────────┴───────────────┼────────────────┴────────────────┘
                                         ▼
                   ┌────────────────────────────────────┐
                   │   SWARM CONSENSUS ENGINE (≥ 80%)   │
                   │  Quant · Catalyst · Security Audit │
                   └────────────────┬───────────────────┘
                                    │ only ≥ 80% confidence
                                    ▼
          MULTI-PLATFORM DISPATCH (Discord · Terminal TUI · Telegram)
                                    │
                                    ▼
                        WALLET TRACKER & POSITION MANAGER
```

### Five Specialist Screening Sub-Agents
1. 🌸 **Meme Robinhood Agent (`#call-meme-robinhood`):** Screens emerging tokens via GMGN OpenAPI & GoPlus Security. Strict default filters: 24h volume ≥ $25k, liquidity ≥ $5k, fee pool ≥ $250, and mandatory smart contract security verification.
2. 🌊 **LP Velocity Agent (`#call-lp-robinhood`):** Scans concentrated liquidity pools on Uniswap V3 via Krystal Cloud API (`ethereum@4663`), filtering for TVL ≥ $10k, 24h volume ≥ $100k, and daily Fee/TVL ≥ 2% for passive yield farmers.
3. 🔮 **NFT Sniper Agent (`#call-nft-robinhood`):** Monitors NFT collections such as Catz NFT on OpenSea REST API v2, detecting floor surges ≥ +10%/1h, volume spikes ≥ 1.5x, and sales velocity ≥ 3 transactions/hour.
4. ☀️ **Alpha Scraper & Sentiment (`#call-alpha-robinhood`):** Harvests 1-hour narrative shifts on Robinhood Chain paired with real-time social sentiment analysis via official X (Twitter) API v2.
5. 🐋 **ETH Whale Tracker (`#call-whale-eth`):** Tracks institutional capital movements on Hyperliquid L1, monitoring ETH perpetual positions ≥ $500k and spot order flow ≥ $50k.

### 3-Layer Swarm Consensus Engine (≥ 80% Confidence Floor)
No individual agent has authority to broadcast signals autonomously. Candidate tokens must undergo a 3-layer quantitative evaluation:
* **Quant & Liquidity Layer (35%):** Depth, slippage, and volume-to-market-cap ratios.
* **Catalyst & Sentiment Layer (35%):** Volume velocity, social sentiment, and smart money accumulation.
* **Security Audit Layer (30%):** Automated anti-honeypot, mintability, blacklist, and ownership checks (*fail-closed* architecture).

Signals scoring below **80%** are immediately discarded. Furthermore, a **Cross-Agent Conflict Veto** cancels BUY recommendations if whale agents detect opposing SHORT or distribution signals on the same underlying asset.

### 9-Lives Risk Engine & Position Manager
Post-execution, the Position Manager safeguards capital:
* **Circuit Breaker:** Emergency shutdown mechanism halting active trading if daily drawdown thresholds are breached.
* **Multi-Tier Take Profit (TP):** Automated profit milestones at +100% (2x) and +200% (3x).
* **Disciplined Stop Loss (SL):** Firm capital preservation boundary at -20%.
* **Dynamic Trailing Stop:** Protects accrued gains as price moves favorably.
* **LP Out-of-Range Alerts:** Real-time warnings when spot price drifts outside active Uniswap V3 liquidity ticks.

---

## 4. Full Strategy Customization & Plain English Compiler

Opencatz AI adapts completely to your personal trading approach:
* **Plain English Strategy Compiler:** Define custom screening criteria using natural language (e.g. *"Only hunt meme tokens held by 3+ smart wallets with liquidity over $15k"*). At startup, the compiler converts this prompt into a validated, sandboxed `.mjs` module.
* **Screening Presets:**
  * *Loosened Default:* Yields ~2x more signals for active traders while maintaining the strict 80% quality threshold.
  * *Standard:* Highly conservative baseline filtering.
  * *Numeric Editor:* Fine-tune numerical thresholds directly during onboarding or via Discord chat.
* **Custom Indicators (`indicators/`):** Dedicated directory for adding custom technical indicators and quantitative formulas.

---

## 5. 💡 Pro-Tip: Run 100% Free with OpenRouter Free Tier

You can operate Opencatz AI around the clock **with zero AI subscription costs ($0/month)**:
1. Create a free account at [OpenRouter.ai](https://openrouter.ai) and generate an API key.
2. Select high-performance free models such as `meta-llama/llama-3.3-70b-instruct:free` or `deepseek/deepseek-r1:free`.
3. Opencatz AI performs all heavy filtering, math scoring, and security audits locally via deterministic code. The LLM is invoked only for high-level reasoning (sentiment analysis and control room chat queries). Operational overhead remains **strictly $0**!

---

## 6. Execution Modes & Multi-Platform Command Center

### Execution Modes
1. **`DRY_RUN` (Default):** Realistic market simulation using real-time Uniswap V3 quotes and gas calculations without capital risk.
2. **`SIGNAL_ONLY`:** Operates purely as an intelligence radar, posting structured call cards with direct swap links for manual execution.
3. **`AUTO_EXECUTE`:** Fully autonomous on-chain trading via Viem and Uniswap V3 when swarm consensus reaches ≥ 80% and risk checks pass.

### Multi-Platform Interfaces
* **🎮 Discord Command Center:** Auto-provisions `🐾 OPENCATZ COMMAND CENTER` category, 6 specialized channels, and 22 slash commands.
* **💻 Terminal TUI (`opencatz terminal`):** Interactive 24-bit TrueColor ANSI dashboard for VPS and headless server management.
* **📱 Telegram Notification Bridge:** High-priority mobile push notifications with interactive callback buttons.
* **🌐 Web Dashboard REST API (Port 3000):** Ready to connect with Next.js dashboards or mobile applications.

---

## 7. 🎟️ Live Deployment in PX Identities Discord (404 Identities Holders)

For traders seeking zero-infrastructure access: **Opencatz AI will be deployed live 24/7 in the PX Identities Discord server!**

Holders of the **404 Identities (Robinhood Chain)** collection receive complimentary access:
* Real-time screening signals across all dedicated channels.
* On-demand smart contract audits directly inside Discord.
* Live whale tracking and daily market summaries with zero server maintenance.

---

## 8. Step-by-Step Installation Guide

### Prerequisites
* **Node.js** version 22.12 or newer ([Download Node.js](https://nodejs.org))
* **Git**
* **Discord Bot Token** from Discord Developer Portal

### Step 1: Clone the Repository
```bash
git clone https://github.com/dizcorvus/opencatz-ai-robinhood-chain.git
cd "Opencatz AI (Robinhood Chain)"
```

### Step 2: Automated 1-Click Setup
* **Windows (PowerShell):**
  ```powershell
  .\\setup.bat
  ```
* **Linux / macOS / VPS:**
  ```bash
  bash setup.sh
  ```

### Step 3: Interactive Onboarding (`opencatz onboard`)
Run the onboarding wizard:
```bash
opencatz onboard
```
Configure execution mode, Discord/Telegram credentials, AI provider, and data API keys (GMGN, Krystal Cloud, OpenSea, GoPlus, Uniswap, X API v2) alongside backup key arrays (`*_BACKUP_KEYS`).

### Step 4: Launch the Swarm
* **Development Mode:** `opencatz run`
* **Terminal TUI Console:** `opencatz terminal`
* **24/7 Production Background Daemon (PM2):** `opencatz deploy`

> ✨ Upon bot initialization, Opencatz **automatically creates the `🐾 OPENCATZ COMMAND CENTER` category, configures all 6 channels, and registers all 22 slash commands**.

### Essential Commands
* `/analyze [address]`: Instant 3-layer security and liquidity audit for any contract.
* `/wallet balance / setup`: Manage burner wallet and inspect on-chain balance.
* `/alert set [token] [target]`: Set automated price alerts with Discord notifications.
* `/journal summary / export`: Review trading performance and export trade logs to CSV.
* `opencatz doctor`: Complete diagnostic check across RPC endpoints, API keys, and sub-agents.
* `opencatz update`: Single-command auto-update (git pull, build, PM2 daemon restart).

---

## 9. 🌟 Open Source & Roadmap

Opencatz AI is **100% Open-Source** under the MIT license. We welcome developers, quants, and crypto researchers to collaborate, build new screening sub-agents, and contribute custom strategy modules.

🔗 **GitHub Repository:** [https://github.com/dizcorvus/opencatz-ai-robinhood-chain](https://github.com/dizcorvus/opencatz-ai-robinhood-chain)

The Robinhood Chain edition represents the foundation of the ecosystem. The team is actively developing **Opencatz AI Multi-Chain Edition** (Solana, Base, Arbitrum, BSC) and a **Premium Swarm Execution Engine**. Stay updated at [opencatz.xyz](https://opencatz.xyz)!
"""

def generate_indonesian_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        hp = section.header.paragraphs[0]
        hp.text = "Opencatz AI — Robinhood Chain Edition | Panduan & Artikel Lengkap"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.name = 'Calibri'
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(148, 163, 184)
        
        fp = section.footer.paragraphs[0]
        fp.text = "Website: https://opencatz.xyz | GitHub: https://github.com/dizcorvus/opencatz-ai-robinhood-chain"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.name = 'Calibri'
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor(148, 163, 184)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("🐾 Mengenal Opencatz AI: Pasukan AI Agent Otonom Pemburu Alpha di Robinhood Chain")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(19)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(10)
    r_sub = p_sub.add_run("Arsitektur Swarm Multi-Agent otonom, konsensus sinyal ketat (≥80%), proteksi risiko on-chain 9-Lives, dan fleksibilitas kustomisasi di Robinhood Chain (EVM L2 #4663).")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(10.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    add_callout(
        doc,
        "\"Chill trades, 9 lives, razor-sharp on-chain alpha.\" — Opencatz AI 🐾⚡\n\n"
        "⚠️ DISCLAIMER AWAL (NFA & DYOR):\n"
        "Seluruh konten artikel ini ditujukan untuk tujuan edukasi dan riset teknologi Web3 semata. Trading aset kripto on-chain memiliki risiko volatilitas tinggi. Selalu lakukan riset mandiri (Do Your Own Research) dan terapkan manajemen risiko yang disiplin. Not Financial Advice.",
        title="💬 QUOTE & RISK REMINDER",
        border_color="E53935",
        bg_color="FFF8F8"
    )

    add_callout(
        doc,
        "🌐 Website Resmi: https://opencatz.xyz\n"
        "📖 Dokumentasi: https://opencatz.xyz/docs\n"
        "💻 Web Terminal Emulator: https://opencatz.xyz/terminal\n"
        "🔗 Repository GitHub: https://github.com/dizcorvus/opencatz-ai-robinhood-chain\n"
        "🎯 Jaringan Target: Robinhood Chain (EVM L2, Chain ID: 4663, Native: ETH, DEX: Uniswap V3)\n"
        "⚡ Fitur Kunci: 3-Layer Swarm Consensus (≥80%), 5 Agen Screening Spesialis, Discord Hub, Terminal TUI, Telegram Bridge, Position Manager & Custom Strategy Compiler.",
        title="📌 RINGKASAN PROYEK",
        border_color="4F46E5",
        bg_color="F8FAFC"
    )

    add_styled_heading(doc, "1. Masalah On-Chain: Kenapa Manual Screening Bikin Capek?", 1)
    add_paragraph_styled(doc, "Buat kamu yang sering mantau pergerakan ekosistem on-chain, pasti tahu rasanya pas sebuah Layer-2 baru rilis. Dalam hitungan menit, ratusan token baru bermunculan, timeline X (Twitter) langsung ramai bahas koin-koin baru, dan chart harga bergerak liar.")
    add_paragraph_styled(doc, "Peluang profit di ekosistem baru memang sangat besar. Tapi realitanya, trading di on-chain itu seperti masuk ke medan perang:")
    
    add_bullet_styled(doc, "Banjir Koin Scam & Rugpull", "Sekitar 90% token baru yang baru deploy berakhir jadi honeypot, likuiditas ditarik dev, atau langsung dihajar sniper bot.")
    add_bullet_styled(doc, "Keterbatasan Fisik Manusia", "Crypto jalan 24/7 tanpa henti. Sering banget momen pergerakan likuiditas besar atau alpha gurih justru meledak jam 3 pagi pas kita lagi istirahat.")
    add_bullet_styled(doc, "Fragmentasi Data", "Mau riset satu token saja kita harus buka DexScreener, GoPlus, GMGN, Twitter/X, dan Krystal secara bersamaan. Capek dan buang waktu.")
    add_bullet_styled(doc, "Jebakan Emosi & FOMO", "Melihat candle hijau panjang membuat trader gatal ingin beli di puncak, lalu panik cutloss saat terjadi koreksi tipis.")

    add_paragraph_styled(doc, "Dari masalah nyata inilah Opencatz AI dibuat. Opencatz AI adalah arsitektur Multi-Agent Swarm Intelligence yang bekerja otonom 24/7 memantau Robinhood Chain, menyaring ribuan token sampah lewat konsensus berlapis, dan mengirimkan hasil analisis siap eksekusi ke Discord, Terminal TUI, atau Telegram kamu.")

    add_styled_heading(doc, "2. Jaringan Spesialis: Robinhood Chain (EVM L2 #4663)", 1)
    add_paragraph_styled(doc, "Opencatz AI tidak dibuat campur aduk dengan bridge lintas-rantai yang rumit. Sistem ini dirancang native dan fokus 100% untuk Robinhood Chain.")

    add_styled_table(doc,
        ["Parameter Jaringan", "Spesifikasi"],
        [
            ["Nama Jaringan", "Robinhood Chain (EVM Layer 2)"],
            ["Chain ID", "4663"],
            ["Native Asset", "ETH"],
            ["Canonical RPC", "https://rpc.mainnet.chain.robinhood.com"],
            ["Block Explorer", "https://robinhoodchain.blockscout.com"],
            ["Primary DEX Venue", "Uniswap V3 Router (Robinhood Chain L2)"],
            ["Karakteristik", "Eksekusi sub-detik dengan gas fee super murah"]
        ]
    )

    add_styled_heading(doc, "3. Bedah Arsitektur & Fitur: Cara Kerja Swarm AI", 1)
    add_paragraph_styled(doc, "Opencatz AI membagi pekerjaan ke 5 sub-agent spesialis, 1 mesin konsensus swarm, dan modul proteksi risiko terintegrasi:")

    add_bullet_styled(doc, "Meme Robinhood Agent (#call-meme-robinhood)", "Berburu koin meme baru via GMGN OpenAPI & GoPlus Security. Filter ketat: volume 24 jam minimal $25k, likuiditas minimal $5k, fee pool minimal $250, dan wajib lolos audit kontrak.")
    add_bullet_styled(doc, "LP Velocity Agent (#call-lp-robinhood)", "Mencari pool likuiditas terkonsentrasi (Uniswap V3) via Krystal Cloud API dengan TVL ≥ $10k, volume 24 jam ≥ $100k, dan rasio harian Fee/TVL ≥ 2% buat passive income.")
    add_bullet_styled(doc, "NFT Sniper Agent (#call-nft-robinhood)", "Memantau floor price dan rarity NFT (seperti Catz NFT) di OpenSea REST API v2 dengan filter lonjakan floor ≥ +10%/jam, volume spike ≥ 1.5x, dan sales velocity minimal 3 transaksi/jam.")
    add_bullet_styled(doc, "Alpha Scraper & Sentimen X (#call-alpha-robinhood)", "Mengikis narasi 1 jam terakhir di Robinhood Chain dipadukan dengan pencarian sentimen real-time via official X API v2.")
    add_bullet_styled(doc, "ETH Whale Tracker (#call-whale-eth)", "Melacak posisi whale di Hyperliquid L1, khususnya posisi perpetual ETH di atas $500k dan order flow spot di atas $50k.")

    add_styled_heading(doc, "3-Layer Swarm Consensus Engine (≥ 80% Quality Floor)", 2)
    add_paragraph_styled(doc, "Setiap kandidat token wajib diuji oleh Swarm Consensus Engine lewat 3 lapisan filter: Kuantitatif & Likuiditas (35%), Katalis & Sentimen (35%), dan Audit Keamanan (30%). Sinyal di bawah 80% langsung dibuang. Dilengkapi juga fitur Cross-Agent Conflict Veto untuk membatalkan sinyal BUY jika agen whale mendeteksi sinyal SHORT pada aset yang sama.")

    add_styled_heading(doc, "Position Manager & 9-Lives Circuit Breaker", 2)
    add_paragraph_styled(doc, "Mengawal posisi terbuka dengan Circuit Breaker otomatis saat pasar anjlok, target Take Profit berjenjang (+100% dan +200%), batas Stop Loss disiplin di -20%, trailing stop dinamis, dan peringatan LP out-of-range.")

    add_styled_heading(doc, "4. Fleksibel & Fully Customizable: Atur Sendiri Strategimu!", 1)
    add_paragraph_styled(doc, "Opencatz AI sangat fleksibel dan dapat dikustomisasi secara menyeluruh:")
    add_bullet_styled(doc, "Custom Prompt Strategy Compiler", "Tulis aturan screening pakai bahasa manusia (contoh: 'Hanya cari token meme yang dipegang minimal 3 smart wallet dan likuiditas di atas $15k'). Prompt ini otomatis dikompilasi menjadi file modul strategi .mjs yang aman pada saat bot pertama kali dijalankan.")
    add_bullet_styled(doc, "Pilihan Preset Hunting", "Pilih Loosened Default (sinyal ~2x lebih aktif dengan tetap menjaga batas kualitas 80%), Standard (konservatif), atau ubah angka batas minimal lewat Numeric Editor.")
    add_bullet_styled(doc, "Custom Indikator Teknis", "Tersedia folder indicators/ untuk menambahkan formula indikator teknikal kustom buatanmu sendiri.")

    add_styled_heading(doc, "5. Tips Hemat: Bisa Dijalankan 100% Gratis Pakai OpenRouter Free Tier", 1)
    add_callout(
        doc,
        "💡 PRO TIP: Kamu bisa menjalankan Opencatz AI 100% GRATIS!\n"
        "1. Ambil API Key gratis di OpenRouter.ai.\n"
        "2. Pilih model gratis seperti meta-llama/llama-3.3-70b-instruct:free atau deepseek/deepseek-r1:free.\n"
        "3. Seluruh proses filtering angka, audit data, dan screening dijalankan menggunakan kalkulasi matematis lokal. AI hanya dipanggil untuk penalaran tingkat tinggi (sentimen tweet & chat room), sehingga biaya operasional bot benar-benar $0!",
        title="💰 ZERO COST RUNNING",
        border_color="16A34A",
        bg_color="F0FDF4"
    )

    add_styled_heading(doc, "6. Tiga Mode Eksekusi & Multi-Platform Command Center", 1)
    add_bullet_styled(doc, "DRY_RUN (Default)", "Simulasi pasar 100% realistis dengan harga live Uniswap V3 tanpa memotong saldo dompet asli.")
    add_bullet_styled(doc, "SIGNAL_ONLY", "Bot bekerja murni sebagai radar intelijen dan mengirim kartu sinyal untuk kamu eksekusi manual.")
    add_bullet_styled(doc, "AUTO_EXECUTE", "Bot mengeksekusi swap otomatis di on-chain via Viem saat sinyal mencapai skor konsensus ≥ 80% dan lolos parameter risiko.")

    add_callout(
        doc,
        "🎟️ LIVE DEPLOY UNTUK HOLDERS 404 IDENTITIES (ROBINHOOD CHAIN)\n"
        "Bagi kamu yang tidak mau repot setup server atau VPS sendiri, Opencatz AI akan segera hadir secara live 24/7 di server Discord PX Identities!\n"
        "Seluruh holder koleksi 404 Identities akan mendapatkan akses eksklusif ke seluruh channel screening, audit on-demand, dan notifikasi whale secara gratis.",
        title="🌟 COMMUNITY ACCESS",
        border_color="0284C7",
        bg_color="F0F9FF"
    )

    add_styled_heading(doc, "7. Panduan Instalasi Langkah Demi Langkah", 1)
    add_paragraph_styled(doc, "Langkah 1: Clone Repositori dari GitHub")
    add_code_block(doc, "git clone https://github.com/dizcorvus/opencatz-ai-robinhood-chain.git\ncd \"Opencatz AI (Robinhood Chain)\"")

    add_paragraph_styled(doc, "Langkah 2: Jalankan Skrip Setup Otomatis")
    add_code_block(doc, "# Windows (PowerShell):\n.\\setup.bat\n\n# Linux / macOS / VPS:\nbash setup.sh")

    add_paragraph_styled(doc, "Langkah 3: Konfigurasi Interaktif (opencatz onboard)")
    add_paragraph_styled(doc, "Wizard akan memandu kamu mengatur mode eksekusi, token Discord/Telegram, kunci AI provider, dan kunci API data (GMGN, Krystal Cloud, OpenSea, GoPlus, Uniswap, X API v2) beserta kunci cadangannya (*_BACKUP_KEYS).")

    add_paragraph_styled(doc, "Langkah 4: Jalankan Bot")
    add_code_block(doc, "# Mode Development:\nopencatz run\n\n# Mode Terminal TUI:\nopencatz terminal\n\n# Mode Produksi 24/7 (PM2 Daemon):\nopencatz deploy")

    add_callout(
        doc,
        "Saat pertama kali bot online di server Discord, Opencatz akan OTOMATIS membuat kategori '🐾 OPENCATZ COMMAND CENTER', menyediakan 6 channel screening, dan mendaftarkan seluruh 22 slash command tanpa perlu dibuat manual.",
        title="✨ ZERO-CONFIG DISCORD",
        border_color="4F46E5",
        bg_color="F8FAFC"
    )

    add_styled_heading(doc, "8. Open Source & Terbuka untuk Kontributor!", 1)
    add_paragraph_styled(doc, "Opencatz AI adalah proyek 100% Open-Source di bawah lisensi MIT. Kami sangat menyambut para developer, trader, dan peneliti on-chain untuk ikut berkontribusi membangun sub-agent baru, modul strategi, atau integrasi data.")
    add_callout(
        doc,
        "🌐 Website Resmi: https://opencatz.xyz\n"
        "🔗 GitHub Repository: https://github.com/dizcorvus/opencatz-ai-robinhood-chain\n"
        "Buka Issue, kirimkan Pull Request, atau berikan Star ⭐️ untuk mendukung proyek ini!",
        title="⭐ OPEN SOURCE & COMMUNITY",
        border_color="16A34A",
        bg_color="F0FDF4"
    )

    add_styled_heading(doc, "9. Bocoran: Versi Multi-Chain & Edisi Premium!", 1)
    add_paragraph_styled(doc, "Arsitektur Opencatz AI di Robinhood Chain ini baru permulaan! Tim saat ini sedang menyiapkan Opencatz AI Multi-Chain Edition (mendukung Solana, Base, Arbitrum, BSC, dll) dan fitur Premium Swarm Execution Engine. Detail lengkap dan akses eksklusifnya akan segera diumumkan melalui portal resmi opencatz.xyz!")

    return doc

def generate_english_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        hp = section.header.paragraphs[0]
        hp.text = "Opencatz AI — Robinhood Chain Edition | Architecture & Setup Guide"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.runs[0].font.name = 'Calibri'
        hp.runs[0].font.size = Pt(8.5)
        hp.runs[0].font.color.rgb = RGBColor(148, 163, 184)
        
        fp = section.footer.paragraphs[0]
        fp.text = "Website: https://opencatz.xyz | GitHub: https://github.com/dizcorvus/opencatz-ai-robinhood-chain"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.runs[0].font.name = 'Calibri'
        fp.runs[0].font.size = Pt(8.5)
        fp.runs[0].font.color.rgb = RGBColor(148, 163, 184)

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("🐾 Inside Opencatz AI: The Autonomous Multi-Agent Trading Swarm on Robinhood Chain")
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(19)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 23, 42)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(10)
    r_sub = p_sub.add_run("An autonomous multi-agent swarm architecture, 3-Layer Consensus (≥80%), on-chain 9-Lives risk protection, and zero-cost deployment on Robinhood Chain L2 (#4663).")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(10.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    add_callout(
        doc,
        "\"Chill trades, 9 lives, razor-sharp on-chain alpha.\" — Opencatz AI 🐾⚡\n\n"
        "⚠️ UPFRONT DISCLAIMER (NFA & DYOR):\n"
        "This article is strictly for educational, research, and technical exploratory purposes. Trading on-chain crypto assets involves substantial risk of capital loss. Always Do Your Own Research (DYOR) and employ strict risk management. Not Financial Advice.",
        title="💬 QUOTE & RISK REMINDER",
        border_color="E53935",
        bg_color="FFF8F8"
    )

    add_callout(
        doc,
        "🌐 Official Website: https://opencatz.xyz\n"
        "📖 Documentation: https://opencatz.xyz/docs\n"
        "💻 Web Terminal Emulator: https://opencatz.xyz/terminal\n"
        "🔗 GitHub Repository: https://github.com/dizcorvus/opencatz-ai-robinhood-chain\n"
        "🎯 Target Network: Robinhood Chain (EVM L2, Chain ID: 4663, Native: ETH, DEX: Uniswap V3)\n"
        "⚡ Key Capabilities: 3-Layer Swarm Consensus (≥80%), 5 Specialist Agents, Discord Hub, Terminal TUI, Telegram Bridge, Position Manager & Plain English Strategy Compiler.",
        title="📌 PROJECT OVERVIEW",
        border_color="4F46E5",
        bg_color="F8FAFC"
    )

    add_styled_heading(doc, "1. The Problem: The Reality of On-Chain Alpha", 1)
    add_paragraph_styled(doc, "Trading on a newly launched Layer-2 network is exhilarating, but anyone who has spent time on-chain knows the reality. Hundreds of new contracts deploy every single hour, token launches flood social feeds, and market liquidity can appear and vanish in seconds.")
    add_paragraph_styled(doc, "The potential upside is massive, but so are the traps:")
    
    add_bullet_styled(doc, "Overwhelming Noise & Honeypots", "Roughly 90% of newly deployed tokens end up as honeypots, sudden developer liquidity pulls (rugpulls), or illiquid pools dominated by sniper bots.")
    add_bullet_styled(doc, "Human Physical Limitations", "Markets operate 24/7. High-conviction alpha and liquidity shifts frequently happen at 3:00 AM while you are away from your desk.")
    add_bullet_styled(doc, "Tool Fragmentation", "Vetting a single token requires juggling DexScreener, GoPlus, GMGN, Twitter/X, and Krystal simultaneously. It is exhausting and slow.")
    add_bullet_styled(doc, "Emotional Trading & FOMO", "Chasing green candles at local tops and panic-selling dips are persistent human pitfalls that erode capital.")

    add_paragraph_styled(doc, "Opencatz AI was built to solve these exact problems: an autonomous multi-agent swarm that continuously monitors Robinhood Chain, applies strict multi-layer mathematical filters, and delivers structured intelligence to Discord, an interactive terminal console, or Telegram.")

    add_styled_heading(doc, "2. Dedicated Focus: Robinhood Chain (EVM L2 #4663)", 1)
    add_paragraph_styled(doc, "Opencatz AI avoids fragile cross-chain bridges and multi-network routing complexity by focusing natively and exclusively on Robinhood Chain.")

    add_styled_table(doc,
        ["Network Parameter", "Specification"],
        [
            ["Network Name", "Robinhood Chain (EVM Layer 2)"],
            ["Chain ID", "4663"],
            ["Native Asset", "ETH (Ethereum)"],
            ["Canonical RPC", "https://rpc.mainnet.chain.robinhood.com"],
            ["Block Explorer", "https://robinhoodchain.blockscout.com"],
            ["Primary DEX Venue", "Uniswap V3 Router (Robinhood Chain L2)"],
            ["Execution Profile", "Sub-second transaction finality with negligible gas fees"]
        ]
    )

    add_styled_heading(doc, "3. System Architecture & Specialist Screening Sub-Agents", 1)
    add_paragraph_styled(doc, "Opencatz AI distributes responsibilities across five specialist screening agents, a central consensus gate, and post-execution risk managers:")

    add_bullet_styled(doc, "Meme Robinhood Agent (#call-meme-robinhood)", "Screens meme tokens via GMGN OpenAPI & GoPlus Security. Enforces 24h volume ≥ $25k, liquidity ≥ $5k, fees ≥ $250, and mandatory security audit passes.")
    add_bullet_styled(doc, "LP Velocity Agent (#call-lp-robinhood)", "Scans concentrated liquidity pools (Uniswap V3) via Krystal Cloud API (ethereum@4663), filtering for TVL ≥ $10k, 24h volume ≥ $100k, and Fee/TVL ≥ 2%.")
    add_bullet_styled(doc, "NFT Sniper Agent (#call-nft-robinhood)", "Tracks NFT collections such as Catz NFT on OpenSea REST v2 with floor surge ≥ +10%/1h, volume spike ≥ 1.5x, and sales velocity ≥ 3/h.")
    add_bullet_styled(doc, "Alpha Scraper & Sentiment (#call-alpha-robinhood)", "Gathers 1-hour narrative shifts paired with official X (Twitter) API v2 real-time sentiment search.")
    add_bullet_styled(doc, "ETH Whale Tracker (#call-whale-eth)", "Monitors Hyperliquid L1 institutional positioning, tracking ETH perps ≥ $500k and spot order flow ≥ $50k.")

    add_styled_heading(doc, "3-Layer Swarm Consensus (≥ 80% Quality Floor)", 2)
    add_paragraph_styled(doc, "Candidates must score ≥ 80% across Quant & Liquidity (35%), Catalyst & Sentiment (35%), and Security Audit (30%). A Cross-Agent Conflict Veto mechanism immediately revokes BUY recommendations if an opposing SHORT signal is active.")

    add_styled_heading(doc, "Position Manager & 9-Lives Risk Shield", 2)
    add_paragraph_styled(doc, "Open positions are protected with automated Take Profit targets (+100% and +200%), a disciplined -20% Stop Loss, dynamic trailing stops, and a circuit breaker for volatile market events.")

    add_styled_heading(doc, "4. Fully Customizable: Tailor Your Own Strategy", 1)
    add_bullet_styled(doc, "Plain English Strategy Compiler", "Define screening rules in plain English (e.g. 'Only hunt meme tokens held by 3+ smart wallets with minimum $15k liquidity'). The system automatically compiles this prompt into a validated, sandboxed .mjs strategy module.")
    add_bullet_styled(doc, "Screening Presets", "Choose Loosened Default (2x more signal volume while keeping the 80% quality floor), Standard (conservative), or fine-tune exact thresholds via the Numeric Editor.")
    add_bullet_styled(doc, "Custom Technical Indicators", "Drop custom indicator formulas into the indicators/ folder.")

    add_styled_heading(doc, "5. Pro-Tip: Run 100% Free with OpenRouter Free Tier", 1)
    add_callout(
        doc,
        "💡 PRO TIP: You can run Opencatz AI completely FREE of charge!\n"
        "1. Create a free account at OpenRouter.ai and generate an API key.\n"
        "2. Select free models like meta-llama/llama-3.3-70b-instruct:free or deepseek/deepseek-r1:free.\n"
        "3. Opencatz AI optimizes token consumption by performing all heavy data filtering and security audits locally via deterministic code. The LLM is only queried for high-value reasoning, keeping running costs literally at $0!",
        title="💰 ZERO COST DEPLOYMENT",
        border_color="16A34A",
        bg_color="F0FDF4"
    )

    add_styled_heading(doc, "6. Execution Modes & Multi-Platform Command Center", 1)
    add_bullet_styled(doc, "DRY_RUN (Default)", "Full simulation mode using real-time quotes and gas calculations with zero financial risk.")
    add_bullet_styled(doc, "SIGNAL_ONLY", "Delivers call cards with direct DEX links for manual execution.")
    add_bullet_styled(doc, "AUTO_EXECUTE", "Autonomous on-chain trading via Viem and Uniswap V3 when consensus is ≥ 80% and risk criteria are met.")

    add_callout(
        doc,
        "🎟️ LIVE DEPLOYMENT FOR 404 IDENTITIES HOLDERS (ROBINHOOD CHAIN)\n"
        "For traders who prefer zero-infrastructure setup, Opencatz AI will be deployed live 24/7 in the PX Identities Discord server soon!\n"
        "All holders of 404 Identities will receive exclusive, zero-config access to real-time screening signals, whale tracking, and on-demand audits.",
        title="🌟 COMMUNITY & HOLDER ACCESS",
        border_color="0284C7",
        bg_color="F0F9FF"
    )

    add_styled_heading(doc, "7. Installation & Setup Guide", 1)
    add_paragraph_styled(doc, "Step 1: Clone the Repository")
    add_code_block(doc, "git clone https://github.com/dizcorvus/opencatz-ai-robinhood-chain.git\ncd \"Opencatz AI (Robinhood Chain)\"")

    add_paragraph_styled(doc, "Step 2: Run the Automated Setup Script")
    add_code_block(doc, "# Windows (PowerShell):\n.\\setup.bat\n\n# Linux / macOS / VPS:\nbash setup.sh")

    add_paragraph_styled(doc, "Step 3: Interactive Onboarding (opencatz onboard)")
    add_paragraph_styled(doc, "Configure execution modes, Discord/Telegram credentials, AI providers, and screening API keys (GMGN, Krystal Cloud, OpenSea, GoPlus, Uniswap, X API v2) along with backup keys (*_BACKUP_KEYS).")

    add_paragraph_styled(doc, "Step 4: Launch the Engine")
    add_code_block(doc, "# Development mode:\nopencatz run\n\n# Terminal TUI console:\nopencatz terminal\n\n# 24/7 Production background daemon (PM2):\nopencatz deploy")

    add_callout(
        doc,
        "Upon joining your Discord server, Opencatz automatically provisions the '🐾 OPENCATZ COMMAND CENTER' category, sets up all 6 channels, and registers all 22 slash commands.",
        title="✨ AUTO-BOOTSTRAP",
        border_color="4F46E5",
        bg_color="F8FAFC"
    )

    add_styled_heading(doc, "8. Open Source & Contributors Welcome!", 1)
    add_paragraph_styled(doc, "Opencatz AI is 100% open-source under the MIT license. We welcome developers, quants, and researchers to contribute new sub-agents, custom strategy modules, and data integrations.")
    add_callout(
        doc,
        "🌐 Official Website: https://opencatz.xyz\n"
        "🔗 GitHub Repository: https://github.com/dizcorvus/opencatz-ai-robinhood-chain\n"
        "Star the repo ⭐️, open issues, or submit Pull Requests to build together!",
        title="⭐ OPEN SOURCE COMMUNITY",
        border_color="16A34A",
        bg_color="F0FDF4"
    )

    add_styled_heading(doc, "9. Teaser: Multichain & Premium Editions on the Horizon!", 1)
    add_paragraph_styled(doc, "The Robinhood Chain edition is just the beginning. The team is actively developing Opencatz AI Multi-Chain Edition (supporting Solana, Base, Arbitrum, BSC, and more) along with a Premium Swarm Execution Engine. Detailed announcements and exclusive access links will drop soon via our official website opencatz.xyz!")

    return doc

def safe_save_doc(doc, target_paths):
    for path in target_paths:
        try:
            doc.save(path)
            print(f"✅ Saved DOCX: {path}")
        except PermissionError:
            alt_path = path.replace(".docx", "_New.docx")
            try:
                doc.save(alt_path)
                print(f"⚠️ File was locked by Word. Saved to: {alt_path}")
            except Exception as e:
                print(f"❌ Failed to save {path}: {e}")
        except Exception as e:
            print(f"❌ Error saving {path}: {e}")

if __name__ == "__main__":
    articles_dir = SCRIPT_DIR
    root_dir = PROJECT_ROOT
    
    # 1. Write Markdown files in articles/ and root
    md_files = [
        (os.path.join(articles_dir, "ARTIKEL_OPENCATZ_AI_INDONESIA.md"), ID_MARKDOWN),
        (os.path.join(articles_dir, "ARTICLE_OPENCATZ_AI_ENGLISH.md"), EN_MARKDOWN),
        (os.path.join(root_dir, "ARTICLE_ID.md"), ID_MARKDOWN),
        (os.path.join(root_dir, "ARTICLE.md"), EN_MARKDOWN),
    ]
    
    for filepath, content in md_files:
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        print(f"✅ Saved Markdown: {filepath}")

    # 2. Build Indonesian DOCX
    print("\n🔨 Generating Indonesian DOCX...")
    doc_id = generate_indonesian_docx()
    safe_save_doc(doc_id, [
        os.path.join(articles_dir, "Opencatz_AI_Robinhood_Chain_Article_ID.docx"),
    ])

    # 3. Build English DOCX
    print("\n🔨 Generating English DOCX...")
    doc_en = generate_english_docx()
    safe_save_doc(doc_en, [
        os.path.join(articles_dir, "Opencatz_AI_Robinhood_Chain_Article_EN.docx"),
    ])

    print("\n🎉 All articles (Markdown & DOCX) successfully generated!")
